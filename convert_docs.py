"""
Script to convert text documents to PDF and DOCX formats.
Run this once to generate the required document formats for the /data folder.
"""
import os
import sys

def create_pdf_from_text(input_path: str, output_path: str) -> None:
    """Convert a text file to PDF using PyMuPDF (fitz)."""
    import fitz  # PyMuPDF

    with open(input_path, "r", encoding="utf-8") as f:
        text = f.read()

    doc = fitz.open()
    lines = text.split("\n")

    # Page layout settings
    page_width, page_height = 595, 842  # A4
    margin = 72
    usable_width = page_width - 2 * margin
    y_pos = margin
    line_height = 14
    title_height = 18

    page = doc.new_page(width=page_width, height=page_height)

    for line in lines:
        stripped = line.strip()

        # Determine font size and style
        if stripped.startswith("Section ") or (stripped and not any(c.islower() for c in stripped[:20]) and len(stripped) < 100):
            font_size = 13
            current_line_height = title_height
        else:
            font_size = 10.5
            current_line_height = line_height

        # Handle empty lines
        if not stripped:
            y_pos += line_height * 0.7
            if y_pos > page_height - margin:
                page = doc.new_page(width=page_width, height=page_height)
                y_pos = margin
            continue

        # Word wrap
        words = stripped.split()
        current_line = ""
        for word in words:
            test_line = f"{current_line} {word}".strip() if current_line else word
            # Approximate character width
            char_width = font_size * 0.5
            if len(test_line) * char_width > usable_width:
                # Write current line
                if y_pos + current_line_height > page_height - margin:
                    page = doc.new_page(width=page_width, height=page_height)
                    y_pos = margin
                page.insert_text(
                    (margin, y_pos),
                    current_line,
                    fontsize=font_size,
                    fontname="helv",
                )
                y_pos += current_line_height
                current_line = word
            else:
                current_line = test_line

        # Write remaining text
        if current_line:
            if y_pos + current_line_height > page_height - margin:
                page = doc.new_page(width=page_width, height=page_height)
                y_pos = margin
            page.insert_text(
                (margin, y_pos),
                current_line,
                fontsize=font_size,
                fontname="helv",
            )
            y_pos += current_line_height

    doc.save(output_path)
    doc.close()
    print(f"  ✓ Created PDF: {output_path} ({os.path.getsize(output_path):,} bytes)")


def create_docx_from_text(input_path: str, output_path: str) -> None:
    """Convert a text file to DOCX using python-docx."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    with open(input_path, "r", encoding="utf-8") as f:
        text = f.read()

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    paragraphs = text.split("\n")
    first_line = True

    for para_text in paragraphs:
        stripped = para_text.strip()

        if not stripped:
            doc.add_paragraph("")
            continue

        if first_line:
            # Document title
            heading = doc.add_heading(stripped, level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            first_line = False
        elif stripped.startswith("Section "):
            # Section heading
            doc.add_heading(stripped, level=1)
        else:
            doc.add_paragraph(stripped)

    doc.save(output_path)
    print(f"  ✓ Created DOCX: {output_path} ({os.path.getsize(output_path):,} bytes)")


def main():
    data_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "data")

    print("=" * 60)
    print("  Document Format Converter")
    print("=" * 60)
    print()

    # Convert coffee_history.txt → coffee_history.pdf
    print("Converting coffee_history.txt → PDF...")
    create_pdf_from_text(
        os.path.join(data_dir, "coffee_history.txt"),
        os.path.join(data_dir, "coffee_history.pdf"),
    )

    # Convert renewable_energy.txt → renewable_energy.pdf
    print("Converting renewable_energy.txt → PDF...")
    create_pdf_from_text(
        os.path.join(data_dir, "renewable_energy.txt"),
        os.path.join(data_dir, "renewable_energy.pdf"),
    )

    # Convert solar_system.txt → solar_system.docx
    print("Converting solar_system.txt → DOCX...")
    create_docx_from_text(
        os.path.join(data_dir, "solar_system.txt"),
        os.path.join(data_dir, "solar_system.docx"),
    )

    print()
    print("✅ All conversions complete!")
    print()
    print("Final document collection:")
    for f in sorted(os.listdir(data_dir)):
        path = os.path.join(data_dir, f)
        if os.path.isfile(path):
            size = os.path.getsize(path)
            ext = os.path.splitext(f)[1].upper()
            print(f"  📄 {f:<40} {ext:<8} {size:>8,} bytes")


if __name__ == "__main__":
    main()
