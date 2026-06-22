"""
Document Loader module for the RAG pipeline.

Supports loading documents in PDF, TXT, and DOCX formats.
Extracts clean text with metadata (source filename, page number, format).
"""

import os
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional


@dataclass
class Document:
    """Represents a loaded document with text content and metadata."""

    text: str
    metadata: dict = field(default_factory=dict)

    @property
    def source(self) -> str:
        return self.metadata.get("source", "unknown")

    @property
    def page(self) -> Optional[int]:
        return self.metadata.get("page")

    @property
    def format(self) -> str:
        return self.metadata.get("format", "unknown")


def load_pdf(file_path: str) -> list[Document]:
    """Load a PDF file and extract text page by page.

    Uses PyMuPDF (fitz) for robust PDF text extraction.
    Strips common headers/footers and page numbers.

    Args:
        file_path: Path to the PDF file.

    Returns:
        List of Document objects, one per page with content.
    """
    import fitz  # PyMuPDF

    documents = []
    filename = os.path.basename(file_path)

    try:
        pdf = fitz.open(file_path)

        for page_num in range(len(pdf)):
            page = pdf[page_num]
            text = page.get_text("text")

            # Clean the extracted text
            text = _clean_text(text)

            if text.strip():
                documents.append(
                    Document(
                        text=text.strip(),
                        metadata={
                            "source": filename,
                            "page": page_num + 1,
                            "format": "pdf",
                            "total_pages": len(pdf),
                        },
                    )
                )

        pdf.close()

    except Exception as e:
        raise RuntimeError(f"Failed to load PDF '{file_path}': {e}") from e

    return documents


def load_txt(file_path: str) -> list[Document]:
    """Load a plain text file and extract text with section detection.

    Splits text by sections (detected via 'Section' headings) to provide
    meaningful page-like metadata.

    Args:
        file_path: Path to the TXT file.

    Returns:
        List of Document objects, one per section.
    """
    filename = os.path.basename(file_path)

    try:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
    except Exception as e:
        raise RuntimeError(f"Failed to load TXT '{file_path}': {e}") from e

    # Split by section headings
    sections = _split_by_sections(text)

    documents = []
    for i, (section_title, section_text) in enumerate(sections):
        cleaned = _clean_text(section_text)
        if cleaned.strip():
            documents.append(
                Document(
                    text=cleaned.strip(),
                    metadata={
                        "source": filename,
                        "page": i + 1,
                        "section": section_title,
                        "format": "txt",
                        "total_pages": len(sections),
                    },
                )
            )

    return documents


def load_docx(file_path: str) -> list[Document]:
    """Load a DOCX file and extract text paragraph by paragraph.

    Uses python-docx for extraction. Groups paragraphs by heading
    sections for meaningful metadata.

    Args:
        file_path: Path to the DOCX file.

    Returns:
        List of Document objects, one per section.
    """
    from docx import Document as DocxDocument

    filename = os.path.basename(file_path)

    try:
        doc = DocxDocument(file_path)
    except Exception as e:
        raise RuntimeError(f"Failed to load DOCX '{file_path}': {e}") from e

    # Group paragraphs by headings
    sections: list[tuple[str, str]] = []
    current_section_title = "Introduction"
    current_text_parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Detect section headings
        if para.style.name.startswith("Heading") or text.startswith("Section "):
            # Save previous section
            if current_text_parts:
                sections.append(
                    (current_section_title, "\n".join(current_text_parts))
                )
            current_section_title = text
            current_text_parts = []
        else:
            current_text_parts.append(text)

    # Don't forget the last section
    if current_text_parts:
        sections.append((current_section_title, "\n".join(current_text_parts)))

    documents = []
    for i, (section_title, section_text) in enumerate(sections):
        cleaned = _clean_text(section_text)
        if cleaned.strip():
            documents.append(
                Document(
                    text=cleaned.strip(),
                    metadata={
                        "source": filename,
                        "page": i + 1,
                        "section": section_title,
                        "format": "docx",
                        "total_pages": len(sections),
                    },
                )
            )

    return documents


def load_documents(directory: str) -> list[Document]:
    """Load all supported documents from a directory.

    Auto-detects file format based on extension and uses the
    appropriate loader. Skips unsupported file types.

    Args:
        directory: Path to the directory containing documents.

    Returns:
        List of all loaded Document objects.
    """
    supported_extensions = {".pdf", ".txt", ".docx"}
    documents = []

    dir_path = Path(directory)
    if not dir_path.exists():
        raise FileNotFoundError(f"Document directory not found: {directory}")

    files = sorted(dir_path.iterdir())

    for file_path in files:
        if not file_path.is_file():
            continue

        ext = file_path.suffix.lower()
        if ext not in supported_extensions:
            continue

        loader = {
            ".pdf": load_pdf,
            ".txt": load_txt,
            ".docx": load_docx,
        }[ext]

        try:
            docs = loader(str(file_path))
            documents.extend(docs)
        except Exception as e:
            print(f"  ⚠ Warning: Failed to load {file_path.name}: {e}")

    return documents


# ---------------------------------------------------------------------------
# Private Helpers
# ---------------------------------------------------------------------------


def _clean_text(text: str) -> str:
    """Clean extracted text by removing artifacts and normalizing whitespace."""
    import re

    # Remove page number patterns (standalone numbers on a line)
    text = re.sub(r"^\s*\d+\s*$", "", text, flags=re.MULTILINE)

    # Remove excessive whitespace while preserving paragraph breaks
    text = re.sub(r"\n{3,}", "\n\n", text)

    # Remove leading/trailing whitespace from each line
    lines = [line.strip() for line in text.split("\n")]
    text = "\n".join(lines)

    return text


def _split_by_sections(text: str) -> list[tuple[str, str]]:
    """Split text into sections based on 'Section X:' headings.

    Args:
        text: The full document text.

    Returns:
        List of (section_title, section_text) tuples.
    """
    import re

    # Split on lines that start with "Section" followed by a number
    pattern = r"^(Section \d+:.*?)$"
    parts = re.split(pattern, text, flags=re.MULTILINE)

    sections = []

    # Handle text before the first section heading
    if parts[0].strip():
        # Check if the first non-empty line looks like a title
        lines = parts[0].strip().split("\n")
        title = lines[0].strip() if lines else "Introduction"
        body = "\n".join(lines[1:]).strip() if len(lines) > 1 else parts[0].strip()
        if body:
            sections.append((title, body))

    # Pair up section headings with their content
    for i in range(1, len(parts), 2):
        heading = parts[i].strip()
        content = parts[i + 1].strip() if i + 1 < len(parts) else ""
        if content:
            sections.append((heading, content))

    # Fallback: if no sections were found, treat entire text as one section
    if not sections:
        sections.append(("Full Document", text.strip()))

    return sections
